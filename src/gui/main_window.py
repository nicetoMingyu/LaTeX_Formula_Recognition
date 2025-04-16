# main_window.py
import os
import logging
import sys
import platform
import tempfile
import time
import numpy as np
from PIL import ImageGrab, Image
from PyQt6.QtWidgets import (
    QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QFileDialog, QProgressBar,
    QTextEdit, QLabel, QMessageBox, QDialog, QSizePolicy,
    QApplication, QListView, QTreeView, QAbstractItemView, QProgressDialog
)
from PyQt6.QtGui import QGuiApplication, QPainter, QPen, QColor, QShortcut, QKeySequence, QFont
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer, QRect, QPoint, QSize
from core.api_client import APIClient
from core.latex_renderer import LatexRenderer
from core.pdf_parser import PDFParser
from config.settings import Config
from gui.api_key_dialog import ApiKeyDialog
from gui.format_dialog import FormatSelectionDialog
from gui.formula_preview_dialog import FormulaPreviewDialog
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Image as RLImage, Paragraph
from reportlab.platypus import Spacer
from reportlab.lib.styles import getSampleStyleSheet
import markdown
import cv2
import base64
from typing import List, Tuple, Optional

class ProcessingThread(QThread):
    progress_updated = pyqtSignal(int)
    task_finished = pyqtSignal(str, bool, str)  # (filename, success, latex)
    processing_done = pyqtSignal()

    def __init__(self, image_paths):
        super().__init__()
        self.image_paths = image_paths
        self.client = APIClient()
        self.logger = logging.getLogger("processing_thread")
        self._cancelled = False
        self.results = []

    def run(self):
        try:
            total = len(self.image_paths)
            for idx, path in enumerate(self.image_paths):
                if self._cancelled:
                    break
                try:
                    formula = self.client.recognize_formula(path)
                    self.task_finished.emit(os.path.basename(path), True, formula)
                    self.results.append(formula)
                    self.progress_updated.emit(int((idx + 1) / total * 100))
                except Exception as e:
                    self.logger.error(f"Process failed {path}: {str(e)}", exc_info=True)
                    self.task_finished.emit(os.path.basename(path), False, "")
            self.processing_done.emit()
        except Exception as e:
            self.logger.error(f"Thread crashed: {str(e)}", exc_info=True)
            self.processing_done.emit()


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("FormulaPro")
        self.setMinimumSize(800, 600)
        self.logger = logging.getLogger("main_window")  # 添加logger
        self._init_ui()
        self._setup_shortcuts()
        self._check_api_key()
        self.api_client = APIClient()  # 添加API客户端
        self.renderer = LatexRenderer()  # 添加渲染器

    def _init_ui(self):
        self.resize(800, 600)

        # UI Components
        self.folder_btn = QPushButton("Select Files/Folder")
        self.screenshot_btn = QPushButton("Screen Capture (Ctrl+Shift+S)")
        self.process_btn = QPushButton("Start Processing")
        self.process_btn.setEnabled(False)
        self.progress_bar = QProgressBar()
        self.editor = QTextEdit()
        self.preview = QLabel()
        self.preview.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.preview.setMinimumSize(400, 200)  # 设置最小尺寸
        self.preview.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)  # 允许水平和垂直方向扩展

        # Layout
        main_layout = QVBoxLayout()
        main_layout.setSpacing(10)  # 设置组件间距
        
        # Top buttons layout
        top_layout = QHBoxLayout()
        top_layout.addWidget(self.folder_btn)
        top_layout.addWidget(self.screenshot_btn)
        top_layout.addWidget(self.process_btn)
        top_layout.addStretch()  # 添加弹性空间
        main_layout.addLayout(top_layout)
        
        # Progress section
        progress_layout = QHBoxLayout()
        progress_layout.addWidget(QLabel("Processing Progress:"))
        progress_layout.addWidget(self.progress_bar)
        progress_layout.addStretch()
        main_layout.addLayout(progress_layout)
        
        # Editor and Preview section
        editor_preview_layout = QHBoxLayout()
        
        # Editor section
        editor_layout = QVBoxLayout()
        editor_layout.addWidget(QLabel("LaTeX Editor:"))
        editor_layout.addWidget(self.editor)
        editor_preview_layout.addLayout(editor_layout, 1)  # 设置拉伸因子为1
        
        # Preview section
        preview_layout = QVBoxLayout()
        preview_layout.addWidget(QLabel("Live Preview:"))
        preview_layout.addWidget(self.preview)
        editor_preview_layout.addLayout(preview_layout, 2)  # 设置拉伸因子为2
        
        main_layout.addLayout(editor_preview_layout)

        # 添加PDF解析按钮
        self.btn_parse_pdf = QPushButton("Parse PDF Formulas")
        self.btn_parse_pdf.clicked.connect(self._parse_pdf)
        main_layout.addWidget(self.btn_parse_pdf)

        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)

    def _setup_shortcuts(self):
        self.folder_btn.clicked.connect(self.select_folder)
        self.process_btn.clicked.connect(self.start_processing)
        self.editor.textChanged.connect(self.update_preview)
        self.screenshot_btn.clicked.connect(self.enter_screenshot_mode)
        self.screenshot_shortcut = QShortcut(QKeySequence("Ctrl+Shift+S"), self)
        self.screenshot_shortcut.activated.connect(self.enter_screenshot_mode)

    def _check_api_key(self):
        dialog = ApiKeyDialog()
        if dialog.exec() != QDialog.DialogCode.Accepted:
            QMessageBox.critical(None, "Error", "Valid API key is required to proceed")
            sys.exit(1)

    # region Screenshot Functionality
    def enter_screenshot_mode(self):
        try:
            # Create a new window for screenshot
            self.capture_overlay = QWidget()
            self.capture_overlay.setWindowFlags(
                Qt.WindowType.FramelessWindowHint |
                Qt.WindowType.WindowStaysOnTopHint |
                Qt.WindowType.Tool
            )
            self.capture_overlay.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
            self.capture_overlay.setStyleSheet("background-color: rgba(0, 0, 0, 50);")

            # Get screen dimensions and set overlay
            screen = QGuiApplication.primaryScreen()
            screen_geometry = screen.geometry()
            self.capture_overlay.setGeometry(screen_geometry)

            # Initialize selection variables
            self.is_selecting = False
            self.selection_start = None
            self.selection_rect = QRect()

            # Add hint label
            self.hint_label = QLabel("Drag mouse to select area | ESC to cancel", self.capture_overlay)
            self.hint_label.setObjectName("screenshot_hint")
            self.hint_label.setStyleSheet("""
                QLabel#screenshot_hint {
                    color: white;
                    background-color: #1a1a1a;
                    padding: 10px;
                    border-radius: 5px;
                    font-size: 16px;
                }
            """)
            self.hint_label.adjustSize()
            self.hint_label.move(
                (screen_geometry.width() - 400) // 2,
                screen_geometry.height() - 80,
            )

            # Hide main window
            self.hide()

            # Bind events to capture overlay
            self.capture_overlay.mousePressEvent = self.selection_start_event
            self.capture_overlay.mouseMoveEvent = self.selection_move_event
            self.capture_overlay.mouseReleaseEvent = self.selection_end_event
            self.capture_overlay.paintEvent = self.paint_selection_rect

            # Show overlay
            self.capture_overlay.show()
            self.capture_overlay.raise_()
            self.capture_overlay.activateWindow()

            # Add escape shortcut to cancel
            self.cancel_shortcut = QShortcut(QKeySequence(Qt.Key.Key_Escape), self.capture_overlay)
            self.cancel_shortcut.activated.connect(self.cancel_screenshot)

        except Exception as e:
            logging.error(f"Failed to enter screenshot mode: {str(e)}", exc_info=True)
            self.cancel_screenshot()

    def selection_start_event(self, event):
        try:
            self.is_selecting = True
            self.selection_start = event.position().toPoint()  # 使用局部坐标
            self.selection_rect = QRect()
            self.capture_overlay.update()
        except Exception as e:
            logging.error(f"Selection start failed: {str(e)}", exc_info=True)

    def selection_move_event(self, event):
        try:
            if not self.is_selecting:
                return

            # 使用局部坐标
            current_pos = event.position().toPoint()
            self.selection_rect = QRect(self.selection_start, current_pos).normalized()
            
            # 更新界面
            self.capture_overlay.update()
            self.update_size_label()
            
        except Exception as e:
            logging.error(f"Selection move failed: {str(e)}", exc_info=True)

    def selection_end_event(self, event):
        try:
            if not self.is_selecting:
                return
                
            self.is_selecting = False
            logging.debug("Selection completed, size: %dx%d",
                          self.selection_rect.width(),
                          self.selection_rect.height())

            if self.selection_rect.width() < 10 or self.selection_rect.height() < 10:
                logging.warning("Selection too small, cancelling")
                self.cancel_screenshot()
                return

            self.capture_overlay.hide()

            QTimer.singleShot(50, self.process_screenshot)

        except Exception as e:
            logging.error(f"Selection end failed: {str(e)}", exc_info=True)
            self.cancel_screenshot()


    def update_size_label(self):
        if not self.selection_rect.isEmpty():
            size_text = f"Selection size: {self.selection_rect.width()}×{self.selection_rect.height()} pixels"
            self.hint_label.setText(size_text)
            self.hint_label.setStyleSheet("""
                QLabel {
                    color: white;
                    font-size: 16px;
                    background: rgba(0,0,0,180);
                    padding: 8px;
                    border-radius: 4px;
                    font-family: Arial;
                    font-weight: bold;
                }
            """)

    def paint_selection_rect(self, event):
        try:
            painter = QPainter(self.capture_overlay)
            painter.setRenderHint(QPainter.RenderHint.Antialiasing)

            # 绘制半透明背景
            painter.fillRect(self.capture_overlay.rect(), QColor(0, 0, 0, 100))

            if hasattr(self, 'selection_rect') and not self.selection_rect.isEmpty():
                # 使用橡皮擦模式清除选择区域
                painter.setCompositionMode(QPainter.CompositionMode.CompositionMode_Clear)
                painter.fillRect(self.selection_rect, Qt.GlobalColor.transparent)
                
                # 恢复正常绘制模式
                painter.setCompositionMode(QPainter.CompositionMode.CompositionMode_SourceOver)

                # 绘制白色实线边框
                painter.setPen(QPen(Qt.GlobalColor.white, 2, Qt.PenStyle.SolidLine))
                painter.drawRect(self.selection_rect)

                # 绘制橙色虚线边框
                pen = QPen(QColor(255, 165, 0), 1, Qt.PenStyle.DashLine)
                pen.setDashPattern([4, 4])
                painter.setPen(pen)
                painter.drawRect(self.selection_rect)

                # 绘制四角手柄
                handle_size = 8
                painter.setPen(Qt.GlobalColor.white)
                painter.setBrush(QColor(255, 165, 0))

                corners = [
                    self.selection_rect.topLeft(),
                    self.selection_rect.topRight(),
                    self.selection_rect.bottomLeft(),
                    self.selection_rect.bottomRight()
                ]

                for corner in corners:
                    painter.drawRect(
                        corner.x() - handle_size // 2,
                        corner.y() - handle_size // 2,
                        handle_size,
                        handle_size
                    )

            painter.end()

        except Exception as e:
            logging.error(f"Paint selection rect failed: {str(e)}", exc_info=True)

    def cancel_screenshot(self):
        try:
            if hasattr(self, 'cancel_shortcut'):
                self.cancel_shortcut.deleteLater()
                del self.cancel_shortcut

            if hasattr(self, 'capture_overlay'):
                self.capture_overlay.close()
                self.capture_overlay = None

            self.show()
            self.activateWindow()

        except Exception as e:
            logging.error(f"Cancel screenshot failed: {str(e)}", exc_info=True)

    def process_screenshot(self):
        try:
            # Validate selection
            if self.selection_rect.isNull() or self.selection_rect.isEmpty():
                logging.error("Invalid screenshot selection")
                self.show()
                return

            # Get screen screenshot
            screen = QGuiApplication.primaryScreen()
            screenshot = screen.grabWindow(0,
                                           self.selection_rect.x(),
                                           self.selection_rect.y(),
                                           self.selection_rect.width(),
                                           self.selection_rect.height())
            if screenshot.isNull():
                raise RuntimeError("Failed to capture screenshot")

            # First show format selection dialog
            format_dialog = FormatSelectionDialog(self)
            if format_dialog.exec() != QDialog.DialogCode.Accepted:
                self.show()
                return

            self.selected_formats = format_dialog.selected_formats()
            if not self.selected_formats:
                QMessageBox.warning(self, "Warning", "Please select at least one format")
                self.show()
                return

            # Get base save path
            base_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Report",
                os.path.expanduser("~/Documents/FormulaReport"),
                "All Files (*)"
            )
            if not base_path:
                logging.info("User cancelled save")
                self.show()
                return

            # Remove extension if present
            self.base_path = os.path.splitext(base_path)[0]

            # Process the screenshot directly
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                if screenshot.save(temp_file.name, "PNG"):
                    self.process_single_image(temp_file.name)
                    self.statusBar().showMessage("Processing screenshot...")
                else:
                    raise RuntimeError("Failed to save temporary screenshot")

        except Exception as e:
            logging.error(f"Screenshot processing failed: {str(e)}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Screenshot processing failed: {str(e)}")
        finally:
            self.show()

    def _get_save_path(self):
        """Get save path independently"""
        self.save_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Report",
            os.path.expanduser("~/Documents/FormulaReport.docx"),
            "Word Documents (*.docx)"
        )

    def process_single_image(self, image_path):
        self.current_thread = ProcessingThread([image_path])
        self.current_thread.progress_updated.connect(self.progress_bar.setValue)
        self.current_thread.task_finished.connect(self.handle_task_result)
        self.current_thread.processing_done.connect(self.save_document)
        self.current_thread.start()

    def select_folder(self):
        try:
            dialog = QFileDialog(self)
            dialog.setWindowTitle("Select Files or Folder")
            dialog.setOption(QFileDialog.Option.DontUseNativeDialog)
            dialog.setFileMode(QFileDialog.FileMode.ExistingFiles)  # Allow selecting multiple files
            dialog.setNameFilter("Image Files (*.png *.jpg *.jpeg);;All Files (*)")
            
            # Get the file view and tree view
            file_view = dialog.findChild(QListView, "listView")
            tree_view = dialog.findChild(QTreeView)
            
            # Enable directory and file selection in the views
            if file_view:
                file_view.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
            if tree_view:
                tree_view.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
                tree_view.setSelectionBehavior(QAbstractItemView.SelectionBehavior.SelectRows)
            
            # Add a Select Folder button
            select_folder_btn = QPushButton("Select Current Folder", dialog)
            select_folder_btn.clicked.connect(lambda: self.handle_folder_selection(dialog))
            
            # Add the button to the dialog's layout
            layout = dialog.layout()
            # Find the button box that contains Open/Cancel buttons
            for i in range(layout.count()):
                widget = layout.itemAt(i).widget()
                if isinstance(widget, QWidget) and widget.layout():
                    # Add our button just above the button box
                    widget.layout().insertWidget(0, select_folder_btn)
                    break
            
            if dialog.exec() == QDialog.DialogCode.Accepted:
                selected_paths = dialog.selectedFiles()
                if not selected_paths:
                    return
                    
                # Process the selected files or folder
                self.process_selected_items(selected_paths)
                    
        except Exception as e:
            logging.error(f"File selection failed: {str(e)}", exc_info=True)
            QMessageBox.critical(self, "Error", f"File selection failed: {str(e)}")

    def handle_folder_selection(self, dialog):
        """Handle the selection of the current folder"""
        try:
            current_dir = dialog.directory().absolutePath()
            dialog.done(QDialog.DialogCode.Accepted)  # Close dialog with accept
            self.process_selected_items([current_dir])  # Process the selected folder
        except Exception as e:
            logging.error(f"Folder selection failed: {str(e)}", exc_info=True)
            QMessageBox.critical(self, "Error", f"Folder selection failed: {str(e)}")

    def process_selected_items(self, paths):
        try:
            # Show format selection dialog first
            format_dialog = FormatSelectionDialog(self)
            if format_dialog.exec() != QDialog.DialogCode.Accepted:
                return

            self.selected_formats = format_dialog.selected_formats()
            if not self.selected_formats:
                QMessageBox.warning(self, "Warning", "Please select at least one format")
                return

            # Get base save path
            base_path, _ = QFileDialog.getSaveFileName(
                self,
                "Save Report",
                os.path.expanduser("~/Documents/FormulaReport"),
                "All Files (*)"
            )
            if not base_path:
                return

            self.base_path = os.path.splitext(base_path)[0]
            
            # Process files
            image_files = []
            
            # Check each path
            for path in paths:
                if os.path.isdir(path):
                    # If it's a directory, get all image files from it
                    for root, _, files in os.walk(path):
                        for file in files:
                            if file.lower().endswith(('.png', '.jpg', '.jpeg')):
                                image_files.append(os.path.join(root, file))
                elif path.lower().endswith(('.png', '.jpg', '.jpeg')):
                    # If it's a file and it's an image, add it directly
                    image_files.append(path)
            
            if not image_files:
                QMessageBox.warning(self, "Warning", "No valid image files found")
                return
                
            # Process all collected image files
            self.current_thread = ProcessingThread(image_files)
            self.current_thread.progress_updated.connect(self.progress_bar.setValue)
            self.current_thread.task_finished.connect(self.handle_task_result)
            self.current_thread.processing_done.connect(self.save_document)
            self.current_thread.start()
            self.statusBar().showMessage("Processing files...")
                    
        except Exception as e:
            logging.error(f"File selection failed: {str(e)}", exc_info=True)
            QMessageBox.critical(self, "Error", f"File selection failed: {str(e)}")

    def start_processing(self):
        # 先选择格式
        format_dialog = FormatSelectionDialog(self)
        if format_dialog.exec() != QDialog.DialogCode.Accepted:
            return

        self.selected_formats = format_dialog.selected_formats()
        if not self.selected_formats:
            QMessageBox.warning(self, "Warning", "Please select at least one format")
            return

        # 获取基础文件名
        base_path, _ = QFileDialog.getSaveFileName(
            self,
            "Save Report",
            os.path.expanduser("~/Documents/FormulaReport"),
            "All Files (*)"
        )
        if not base_path:
            return

        self.base_path = os.path.splitext(base_path)[0]  # 去除扩展名
        self.progress_bar.setValue(0)
        self.editor.clear()
        self.preview.clear()

        try:
            images = [
                os.path.join(self.image_folder, f)
                for f in os.listdir(self.image_folder)
                if f.lower().endswith((".png", ".jpg", ".jpeg"))
                and os.path.isfile(os.path.join(self.image_folder, f))
            ]
            if not images:
                QMessageBox.warning(self, "Warning", "No valid images found!")
                return
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Folder error: {str(e)}")
            return

        self.current_thread = ProcessingThread(images)
        self.current_thread.progress_updated.connect(self.progress_bar.setValue)
        self.current_thread.task_finished.connect(self.handle_task_result)
        self.current_thread.processing_done.connect(self.save_document)
        self.current_thread.start()

        QTimer.singleShot(Config.TIMEOUT * 1000, self.check_thread_status)

    def check_thread_status(self):
        if self.current_thread and self.current_thread.isRunning():
            QMessageBox.warning(self, "Timeout", "Processing timeout occurred!")
            self.current_thread.cancel()

    def handle_task_result(self, filename, success, latex):
        status = "✓" if success else "✗"
        self.statusBar().showMessage(f"{status} Processed {filename}")
        if success:
            self.editor.setPlainText(latex)
            self.update_preview()

    def update_preview(self):
        """Update formula preview"""
        code = self.editor.toPlainText()
        try:
            if code.strip():
                pixmap = self.renderer.render_to_qpixmap(code)
                if pixmap and not pixmap.isNull():
                    # 计算预览框的大小
                    preview_size = self.preview.size()
                    # 保持宽高比例缩放图片
                    scaled_pixmap = pixmap.scaled(
                        preview_size,
                        Qt.AspectRatioMode.KeepAspectRatio,
                        Qt.TransformationMode.SmoothTransformation
                    )
                    self.preview.setPixmap(scaled_pixmap)
                    QApplication.processEvents()  # 确保预览立即更新
        except Exception as e:
            self.logger.error(f"Preview error: {str(e)}", exc_info=True)

    def save_document(self):
        try:
            if not hasattr(self, 'base_path'):
                raise ValueError("Save path not selected")

            # Generate files for each format
            success_formats = []
            for fmt in self.selected_formats:
                try:
                    if fmt == "docx":
                        self._save_docx()
                        success_formats.append("docx (1 file)")
                    elif fmt == "pdf":
                        self._save_pdf()
                        success_formats.append("pdf (1 file)")
                    elif fmt == "tex":
                        self._save_tex()
                        success_formats.append("tex (1 file)")
                    elif fmt == "md":
                        self._save_md()
                        success_formats.append("md (1 file)")
                    elif fmt == "svg":
                        self._save_svg()
                        success_formats.append(f"svg ({len(self.current_thread.results)} files)")
                    elif fmt == "png":
                        self._save_png()
                        success_formats.append(f"png ({len(self.current_thread.results)} files)")
                except Exception as e:
                    logging.error(f"Failed to save {fmt.upper()}: {str(e)}", exc_info=True)

            if success_formats:
                QMessageBox.information(
                    self,
                    "Success",
                    f"Saved formats:\n{', '.join(success_formats)}\nBase path: {self.base_path}"
                )
            else:
                QMessageBox.warning(self, "Error", "All format saves failed, please check logs")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Save failed: {str(e)}")
            logging.error(f"Document save error: {str(e)}", exc_info=True)

    def _save_docx(self):
        """保存Word文档（原有逻辑）"""
        from docx import Document
        from docx.shared import Inches
        from docx.oxml.ns import qn
        import tempfile

        try:
            doc = Document()  # 确保这行缩进正确
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            doc.add_heading('Formula Recognition Report', 0)

            for idx, formula in enumerate(self.current_thread.results):
                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_path = os.path.join(temp_dir, f"temp_{idx}.png")
                    try:
                        pixmap = self.renderer.render_to_qpixmap(formula)
                        if pixmap.save(temp_path, "PNG"):
                            doc.add_heading(f"Formula {idx + 1}", level=2)
                            doc.add_paragraph(f"LaTeX Code:\n{formula}")
                            doc.add_picture(temp_path, width=Inches(5))
                        else:
                            raise RuntimeError("Failed to save image")
                    except Exception as e:
                        logging.error(f"Failed to save formula {idx}: {str(e)}")
                        doc.add_paragraph(f"Formula {idx + 1} render failed: {str(e)}")

            doc.save(f"{self.base_path}.docx")
        except Exception as e:
            raise RuntimeError(f"Failed to save DOCX: {str(e)}")

    def _save_pdf(self):
        """Generate PDF identical to Word (Fixed version)"""
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        import tempfile
        import shutil

        try:
            # Create persistent temporary directory
            temp_dir = tempfile.mkdtemp(prefix="formulapro_")
            logging.info(f"PDF temp directory created at: {temp_dir}")

            # Create PDF document framework
            doc = SimpleDocTemplate(
                f"{self.base_path}.pdf",
                pagesize=A4,
                leftMargin=20 * mm,
                rightMargin=20 * mm,
                topMargin=20 * mm,
                bottomMargin=20 * mm
            )
            styles = getSampleStyleSheet()
            elements = []

            # Add main title
            elements.append(Paragraph(
                "<font size=18><b>Formula Recognition Report</b></font>",
                styles["Title"]
            ))
            elements.append(Spacer(1, 15 * mm))

            for idx, formula in enumerate(self.current_thread.results):
                # Generate image file path
                img_path = os.path.join(temp_dir, f"formula_{idx}.png")

                # Ensure rendering and saving operations are successful
                pixmap = self.renderer.render_to_qpixmap(formula)
                if pixmap.isNull():
                    raise RuntimeError(f"Formula {idx + 1} rendering failed")

                # Save image using absolute path
                if not pixmap.save(img_path, "PNG", quality=100):
                    raise RuntimeError(f"Cannot save formula image: {img_path}")

                # Verify file exists
                if not os.path.exists(img_path):
                    raise FileNotFoundError(f"Image file not generated: {img_path}")

                # Add PDF elements
                elements.append(Paragraph(
                    f"<b>Formula {idx + 1}</b>",
                    styles["Heading2"]
                ))

                elements.append(Paragraph(
                    f"LaTeX Code: <font face='Courier'>{formula}</font>",
                    styles["BodyText"]
                ))

                # Load image using absolute path
                full_img_path = os.path.abspath(img_path)
                img = Image(full_img_path, width=150 * mm, height=40 * mm)
                elements.append(img)

                elements.append(Spacer(1, 10 * mm))

            # Build PDF document
            doc.build(elements)
            logging.info("PDF generated successfully")

        except Exception as e:
            logging.error(f"PDF save failed: {str(e)}", exc_info=True)
            raise RuntimeError(f"PDF generation failed: {str(e)}")
        finally:
            # Clean up temporary directory
            if os.path.exists(temp_dir):
                try:
                    shutil.rmtree(temp_dir)
                    logging.info(f"Cleaned up temporary directory: {temp_dir}")
                except Exception as e:
                    logging.warning(f"Failed to clean up temporary directory: {str(e)}")

    def _save_tex(self):
        """生成可直接编译的LaTeX文档"""
        tex_content = """\\documentclass{article}
    \\usepackage{amsmath}
    \\begin{document}
    \\title{Formula Report}
    \\maketitle
    """

        for idx, formula in enumerate(self.current_thread.results):
            # 清理公式环境符号
            clean_formula = formula.replace("\\begin{equation}", "") \
                .replace("\\end{equation}", "") \
                .strip("$")

            tex_content += f"\\section*{{Formula {idx + 1}}}\n"
            tex_content += "\\begin{align*}\n"
            tex_content += f"{clean_formula}\n"
            tex_content += "\\end{align*}\n\n"

        tex_content += "\\end{document}"

        with open(f"{self.base_path}.tex", "w") as f:
            f.write(tex_content)

    def _save_md(self):
        """生成可直接显示数学公式的Markdown"""
        md_content = "# Formula Recognition Report\n\n"

        for idx, formula in enumerate(self.current_thread.results):
            # 直接写入LaTeX公式（不带$符号）
            clean_formula = formula.strip("$")

            md_content += f"## Formula {idx + 1}\n\n"
            md_content += "**LaTeX Code:**\n"
            md_content += "```math\n"  # 使用代码块包裹
            md_content += f"{clean_formula}\n"
            md_content += "```\n\n"
            md_content += f"**Rendered Formula:**\n$$\n{clean_formula}\n$$\n\n"  # 内联公式

        with open(f"{self.base_path}.md", "w") as f:
            f.write(md_content)

    def _save_svg(self):
        """Save formulas as SVG files"""
        try:
            from PyQt6.QtSvg import QSvgGenerator
            from PyQt6.QtCore import QSize

            for idx, formula in enumerate(self.current_thread.results):
                # Create SVG generator for each formula
                generator = QSvgGenerator()
                svg_path = f"{self.base_path}_{idx + 1}.svg"
                generator.setFileName(svg_path)
                
                # Get formula pixmap to determine size
                pixmap = self.renderer.render_to_qpixmap(formula)
                generator.setSize(pixmap.size())
                generator.setViewBox(QRect(0, 0, pixmap.width(), pixmap.height()))

                # Create painter
                painter = QPainter()
                painter.begin(generator)

                # Set up rendering
                painter.setRenderHint(QPainter.RenderHint.Antialiasing)
                painter.setRenderHint(QPainter.RenderHint.SmoothPixmapTransform)

                # Draw formula
                painter.drawPixmap(0, 0, pixmap)

                painter.end()
                logging.info(f"SVG file {idx + 1} saved successfully: {svg_path}")

        except Exception as e:
            raise RuntimeError(f"Failed to save SVG: {str(e)}")

    def _save_png(self):
        """Save formulas as PNG files"""
        try:
            for idx, formula in enumerate(self.current_thread.results):
                # Get formula pixmap
                pixmap = self.renderer.render_to_qpixmap(formula)
                
                # Save as PNG with index
                png_path = f"{self.base_path}_{idx + 1}.png"
                if pixmap.save(png_path, "PNG"):
                    logging.info(f"PNG file {idx + 1} saved successfully: {png_path}")
                else:
                    raise RuntimeError(f"Failed to save PNG file {idx + 1}")

        except Exception as e:
            raise RuntimeError(f"Failed to save PNG: {str(e)}")

    def _parse_pdf(self):
        """Parse PDF file and extract formulas"""
        try:
            # 选择PDF文件
            pdf_path, _ = QFileDialog.getOpenFileName(
                self,
                "Select PDF File",
                "",
                "PDF Files (*.pdf)"
            )
            
            if not pdf_path:
                return
                
            # 显示进度对话框
            progress = QProgressDialog("Parsing PDF...", None, 0, 100, self)
            progress.setWindowModality(Qt.WindowModality.WindowModal)
            progress.setMinimumDuration(0)
            progress.setAutoClose(False)
            progress.setAutoReset(False)
            progress.show()
            
            try:
                # 解析PDF
                parser = PDFParser()
                formulas = parser.extract_formulas(pdf_path)
            finally:
                progress.close()
            
            if not formulas:
                QMessageBox.warning(self, "Warning", "No formulas found in PDF")
                return
                
            # 显示公式预览对话框
            preview_dialog = FormulaPreviewDialog(formulas, self)
            if preview_dialog.exec() != QDialog.DialogCode.Accepted:
                return
                
            selected_formulas = preview_dialog.get_selected_formulas()
            if not selected_formulas:
                QMessageBox.warning(self, "Warning", "Please select at least one formula")
                return
                
            # 显示格式选择对话框
            format_dialog = FormatSelectionDialog(self)
            if format_dialog.exec() != QDialog.DialogCode.Accepted:
                return
                
            # 获取选择的格式
            selected_formats = format_dialog.selected_formats()
            if not selected_formats:
                QMessageBox.warning(self, "Warning", "Please select at least one output format")
                return
                
            # 选择保存目录
            save_dir = QFileDialog.getExistingDirectory(
                self,
                "Select Save Directory",
                ""
            )
            
            if not save_dir:
                return
                
            # 处理选中的公式
            self._process_formulas(selected_formulas, selected_formats, save_dir)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error processing PDF: {str(e)}")
            self.logger.error(f"Error processing PDF: {str(e)}")
            
    def _process_formulas(self, formulas: List[Tuple[np.ndarray, Tuple[float, float, float, float], float]], 
                         formats: List[str], save_dir: str):
        """Process selected formulas"""
        try:
            self.progress_bar.setValue(0)
            self.progress_bar.setMaximum(len(formulas))
            
            results = []  # 存储识别结果
            
            for i, formula_data in enumerate(formulas):
                self.progress_bar.setValue(i)
                self.statusBar().showMessage(f"Processing formula {i+1}/{len(formulas)}...")
                QApplication.processEvents()  # 处理事件循环，确保UI更新
                
                # 解包公式数据
                if len(formula_data) == 2:
                    formula, position = formula_data
                    confidence = 1.0
                else:
                    formula, position, confidence = formula_data
                
                # 识别公式
                latex = self._recognize_formula(formula)
                if latex:
                    results.append(latex)
                    # 更新编辑器和预览
                    self.editor.setPlainText(latex)
                    self.update_preview()
                    QApplication.processEvents()  # 再次处理事件循环，确保预览更新
            
            if not results:
                QMessageBox.warning(self, "Warning", "No formulas were successfully processed")
                return
                
            # 保存所有格式的文件
            self._save_all_formats(results, formats, save_dir)
            
            self.progress_bar.setValue(len(formulas))
            QMessageBox.information(self, "Success", "Formulas processed successfully!")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Error processing formulas: {str(e)}")
            self.logger.error(f"Error processing formulas: {str(e)}")
            
    def _save_all_formats(self, formulas: List[str], formats: List[str], save_dir: str):
        """Save all formulas in selected formats"""
        try:
            for fmt in formats:
                if fmt in ['png', 'svg']:
                    # 单独保存每个公式
                    for i, formula in enumerate(formulas):
                        output_path = os.path.join(save_dir, f"formula_{i+1}.{fmt}")
                        try:
                            pixmap = self.renderer.render_to_qpixmap(formula)
                            if pixmap:
                                if fmt == 'svg':
                                    # 使用QSvgGenerator保存SVG
                                    from PyQt6.QtSvg import QSvgGenerator
                                    from PyQt6.QtCore import QRectF
                                    generator = QSvgGenerator()
                                    generator.setFileName(output_path)
                                    generator.setSize(pixmap.size())
                                    generator.setViewBox(QRectF(0, 0, pixmap.width(), pixmap.height()))
                                    
                                    painter = QPainter()
                                    painter.begin(generator)
                                    painter.drawPixmap(0, 0, pixmap)
                                    painter.end()
                                    self.logger.info(f"SVG file {i+1} saved successfully: {output_path}")
                                else:
                                    pixmap.save(output_path, "PNG")
                                    self.logger.info(f"PNG file {i+1} saved successfully: {output_path}")
                        except Exception as e:
                            self.logger.error(f"Failed to save {fmt.upper()} file {i+1}: {str(e)}")
                else:
                    # 合并保存到一个文件
                    output_path = os.path.join(save_dir, f"FormulaReport.{fmt}")
                    if fmt == 'docx':
                        self._save_combined_docx(formulas, output_path)
                    elif fmt == 'pdf':
                        self._save_combined_pdf(formulas, output_path)
                    elif fmt == 'tex':
                        self._save_combined_tex(formulas, output_path)
                    elif fmt == 'md':
                        self._save_combined_md(formulas, output_path)
                        
        except Exception as e:
            self.logger.error(f"Error saving files: {str(e)}")
            raise
            
    def _save_combined_docx(self, formulas: List[str], output_path: str):
        """Save all formulas in one Word document"""
        from docx import Document
        from docx.shared import Inches
        from docx.oxml.ns import qn
        import tempfile

        try:
            doc = Document()
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

            doc.add_heading('Formula Recognition Report', 0)

            for idx, formula in enumerate(formulas):
                with tempfile.TemporaryDirectory() as temp_dir:
                    temp_path = os.path.join(temp_dir, f"temp_{idx}.png")
                    try:
                        pixmap = self.renderer.render_to_qpixmap(formula)
                        if pixmap.save(temp_path, "PNG"):
                            doc.add_heading(f"Formula {idx + 1}", level=2)
                            doc.add_paragraph(f"LaTeX Code:\n{formula}")
                            doc.add_picture(temp_path, width=Inches(5))
                        else:
                            raise RuntimeError("Failed to save image")
                    except Exception as e:
                        self.logger.error(f"Failed to save formula {idx}: {str(e)}")
                        doc.add_paragraph(f"Formula {idx + 1} render failed: {str(e)}")

            doc.save(output_path)
            self.logger.info(f"Word document saved successfully: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to save DOCX: {str(e)}")
            raise
            
    def _save_combined_pdf(self, formulas: List[str], output_path: str):
        """Save all formulas in one PDF document"""
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Image, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import mm
        import tempfile

        try:
            # Create temporary directory
            with tempfile.TemporaryDirectory() as temp_dir:
                # Create PDF document
                doc = SimpleDocTemplate(
                    output_path,
                    pagesize=A4,
                    leftMargin=20*mm,
                    rightMargin=20*mm,
                    topMargin=20*mm,
                    bottomMargin=20*mm
                )
                
                styles = getSampleStyleSheet()
                elements = []

                # Add title
                elements.append(Paragraph(
                    "<font size=18><b>Formula Recognition Report</b></font>",
                    styles["Title"]
                ))
                elements.append(Spacer(1, 15*mm))

                # Add each formula
                for idx, formula in enumerate(formulas):
                    # Save formula image
                    img_path = os.path.join(temp_dir, f"formula_{idx}.png")
                    pixmap = self.renderer.render_to_qpixmap(formula)
                    if not pixmap.save(img_path, "PNG"):
                        raise RuntimeError(f"Failed to save formula image {idx}")

                    # Add formula section
                    elements.append(Paragraph(
                        f"<b>Formula {idx + 1}</b>",
                        styles["Heading2"]
                    ))
                    elements.append(Paragraph(
                        f"LaTeX Code: <font face='Courier'>{formula}</font>",
                        styles["BodyText"]
                    ))
                    elements.append(Image(img_path, width=150*mm, height=40*mm))
                    elements.append(Spacer(1, 10*mm))

                # Build PDF
                doc.build(elements)
                self.logger.info(f"PDF document saved successfully: {output_path}")
                
        except Exception as e:
            self.logger.error(f"Failed to save PDF: {str(e)}")
            raise
            
    def _save_combined_tex(self, formulas: List[str], output_path: str):
        """Save all formulas in one LaTeX document"""
        try:
            tex_content = """\\documentclass{article}
\\usepackage{amsmath}
\\begin{document}
\\title{Formula Recognition Report}
\\maketitle
"""

            for idx, formula in enumerate(formulas):
                # Clean formula
                clean_formula = formula.replace("\\begin{equation}", "") \
                    .replace("\\end{equation}", "") \
                    .strip("$")

                tex_content += f"\\section*{{Formula {idx + 1}}}\n"
                tex_content += "\\begin{align*}\n"
                tex_content += f"{clean_formula}\n"
                tex_content += "\\end{align*}\n\n"

            tex_content += "\\end{document}"

            with open(output_path, "w", encoding='utf-8') as f:
                f.write(tex_content)
                
            self.logger.info(f"LaTeX document saved successfully: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to save TEX: {str(e)}")
            raise
            
    def _save_combined_md(self, formulas: List[str], output_path: str):
        """Save all formulas in one Markdown document"""
        try:
            md_content = "# Formula Recognition Report\n\n"

            for idx, formula in enumerate(formulas):
                clean_formula = formula.strip("$")
                
                md_content += f"## Formula {idx + 1}\n\n"
                md_content += "**LaTeX Code:**\n"
                md_content += "```math\n"
                md_content += f"{clean_formula}\n"
                md_content += "```\n\n"
                md_content += f"**Rendered Formula:**\n$$\n{clean_formula}\n$$\n\n"

            with open(output_path, "w", encoding='utf-8') as f:
                f.write(md_content)
                
            self.logger.info(f"Markdown document saved successfully: {output_path}")
            
        except Exception as e:
            self.logger.error(f"Failed to save MD: {str(e)}")
            raise

    def _recognize_formula(self, formula: np.ndarray) -> Optional[str]:
        """Recognize formula using API"""
        try:
            # 创建临时文件来保存图片
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp_file:
                cv2.imwrite(temp_file.name, formula)
                
                # 调用API
                response = self.api_client.recognize_formula(temp_file.name)
                
                # 删除临时文件
                os.unlink(temp_file.name)
                
                # 检查响应类型
                if isinstance(response, dict):
                    return response.get('latex')
                elif isinstance(response, str):
                    return response
                else:
                    self.logger.error(f"Unexpected response type: {type(response)}")
                    return None
            
        except Exception as e:
            self.logger.error(f"Error recognizing formula: {str(e)}")
            return None

    def closeEvent(self, event):
        """Handle window close event"""
        try:
            # 取消所有正在进行的操作
            if hasattr(self, 'current_thread') and self.current_thread and self.current_thread.isRunning():
                self.current_thread.cancel()
                self.current_thread.wait(3000)
            
            # 关闭所有子窗口
            for widget in QApplication.topLevelWidgets():
                if widget != self and isinstance(widget, QDialog):
                    widget.close()
                    
            event.accept()
        except Exception as e:
            self.logger.error(f"Error in closeEvent: {str(e)}")
            event.accept()